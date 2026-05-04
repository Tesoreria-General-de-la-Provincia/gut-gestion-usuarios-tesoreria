from flask import Flask, render_template, request, jsonify
from dotenv import load_dotenv
import logging
import os
import re
import socket
import subprocess

load_dotenv(os.path.join(os.path.dirname(__file__), "..", ".env"))

logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s")
log = logging.getLogger("gut")

UPASS_RE = re.compile(r"^[A-Za-z0-9]{8}$")
EMAIL_RE = re.compile(r"^[^@\s]+@[^@\s]+\.[^@\s]+$")

app = Flask(__name__)

# ─────────────────────────────────────────
#  MODO: "sqlite" para desarrollo local
#         "sqlserver" para producción
#  Se controla con la env var GUT_MODO
# ─────────────────────────────────────────
MODO = os.getenv("GUT_MODO", "sqlserver")

# ── Configuración SQL Server (producción) ─
# Valores leídos desde env vars (.env en local, env_file en docker)
DB_CONFIG = {
    "server":   os.getenv("GUT_DB_SERVER",   ""),
    "database": os.getenv("GUT_DB_NAME",     ""),
    "username": os.getenv("GUT_DB_USER",     ""),
    "password": os.getenv("GUT_DB_PASSWORD", ""),
}

if MODO == "sqlserver" and not all(DB_CONFIG.values()):
    raise RuntimeError(
        "Faltan variables de entorno para SQL Server: "
        "GUT_DB_SERVER, GUT_DB_NAME, GUT_DB_USER, GUT_DB_PASSWORD"
    )

def _net_check():
    host = os.getenv("GUT_NET_CHECK_HOST", "").strip()
    port = os.getenv("GUT_NET_CHECK_PORT", "").strip()
    timeout = float(os.getenv("GUT_NET_CHECK_TIMEOUT", "5"))

    if not host:
        log.info("Net check: deshabilitado (definí GUT_NET_CHECK_HOST en .env para activar)")
        return

    if port:
        try:
            port_int = int(port)
        except ValueError:
            log.warning("Net check: GUT_NET_CHECK_PORT='%s' no es un número, se omite", port)
            return
        log.info("Net check (TCP): probando %s:%s (timeout=%ss)...", host, port_int, timeout)
        try:
            with socket.create_connection((host, port_int), timeout=timeout):
                log.info("Net check OK: %s:%s alcanzable desde el contenedor", host, port_int)
        except socket.timeout:
            log.error("Net check FAIL: timeout conectando a %s:%s — el contenedor no rutea o hay firewall", host, port_int)
        except OSError as e:
            log.error("Net check FAIL: %s:%s no alcanzable (%s)", host, port_int, e)
        return

    log.info("Net check (ICMP): ping a %s (timeout=%ss)...", host, timeout)
    try:
        res = subprocess.run(
            ["ping", "-c", "2", "-W", str(int(timeout)), host],
            capture_output=True, text=True, timeout=timeout + 5,
        )
        if res.returncode == 0:
            log.info("Net check OK: %s responde a ping desde el contenedor", host)
        else:
            log.error("Net check FAIL: %s NO responde a ping (rc=%s). stderr=%s", host, res.returncode, res.stderr.strip() or res.stdout.strip())
    except FileNotFoundError:
        log.error("Net check FAIL: el binario 'ping' no está instalado en la imagen")
    except subprocess.TimeoutExpired:
        log.error("Net check FAIL: ping a %s expiró (timeout)", host)

if MODO == "sqlserver":
    _net_check()

TABLE_NAME   = "TW_Usuarios"
SEARCH_FIELD = "CUIT"

READ_FIELDS = [
    "CUIT",
    "IDProvincia",
    "Descripcion",
    "Domicilio",
    "CtaBco",
]
EDIT_FIELDS = [
    "Fech_created",
    "UPass",
    "Email",
]
# ─────────────────────────────────────────

def get_connection():
    if MODO == "sqlite":
        import sqlite3
        db_path = os.path.join(os.path.dirname(__file__), "gut_dev.db")
        conn = sqlite3.connect(db_path)
        conn.row_factory = sqlite3.Row
        return conn, "sqlite"
    else:
        import pyodbc
        conn_str = (
            f"DRIVER={{ODBC Driver 17 for SQL Server}};"
            f"SERVER={DB_CONFIG['server']};"
            f"DATABASE={DB_CONFIG['database']};"
            f"UID={DB_CONFIG['username']};"
            f"PWD={DB_CONFIG['password']};"
            f"TrustServerCertificate=yes;"
            f"Connection Timeout=5;"
        )
        log.info("Conectando a SQL Server: server=%s db=%s user=%s", DB_CONFIG['server'], DB_CONFIG['database'], DB_CONFIG['username'])
        return pyodbc.connect(conn_str, timeout=5), "sqlserver"


def init_sqlite():
    import sqlite3
    db_path = os.path.join(os.path.dirname(__file__), "gut_dev.db")
    conn = sqlite3.connect(db_path)
    cur  = conn.cursor()
    cur.execute(f"""
        CREATE TABLE IF NOT EXISTS {TABLE_NAME} (
            IDUsuarios   INTEGER PRIMARY KEY,
            IDProvincia  INTEGER,
            CUIT         TEXT NOT NULL UNIQUE,
            Descripcion  TEXT,
            Fech_created TEXT,
            UPass        TEXT,
            Domicilio    TEXT,
            CtaBco       TEXT,
            Email        TEXT,
            auditoria    TEXT
        )
    """)
    cur.execute(f"SELECT COUNT(*) FROM {TABLE_NAME}")
    if cur.fetchone()[0] == 0:
        datos = [
            (1, 1, "20123456780", "GARCIA JUAN CARLOS",     "2023-05-10", "GARCIA01", "AV. MITRE 450",     "0010203040", "jgarcia@gmail.com",      "sistema"),
            (2, 1, "27277781919", "RODRIGUEZ MARIA LAURA",  "2022-11-20", "RODRIG22", "SAN MARTIN 1200",   "0011223344", "mrodriguez@hotmail.com", "sistema"),
            (3, 2, "30654321098", "EMPRESA SRL",            "2024-01-15", "EMPRE30",  "BELGRANO 890 OF 3", "0099887766", None,                     "sistema"),
            (4, 1, "20987654321", "LOPEZ ROBERTO DANIEL",   "2021-08-03", "LOPEZ99",  "LAS HERAS 340",     "0055443322", "rlopez@yahoo.com",       "sistema"),
            (5, 3, "27345678901", "FERNANDEZ ANA BEATRIZ",  "2023-12-01", "FERNA27",  "9 DE JULIO 780",    "0012345678", None,                     "sistema"),
        ]
        cur.executemany(
            f"INSERT INTO {TABLE_NAME} (IDUsuarios,IDProvincia,CUIT,Descripcion,Fech_created,UPass,Domicilio,CtaBco,Email,auditoria) VALUES (?,?,?,?,?,?,?,?,?,?)",
            datos
        )
    conn.commit()
    conn.close()


def limpiar_cuit(valor):
    return valor.strip().replace("-", "").replace(" ", "")


@app.route("/")
def index():
    return render_template("index.html",
                           search_field=SEARCH_FIELD,
                           read_fields=READ_FIELDS,
                           edit_fields=EDIT_FIELDS,
                           table_name=TABLE_NAME,
                           modo=MODO)


@app.route("/buscar", methods=["POST"])
def buscar():
    raw   = request.json.get("valor", "")
    valor = limpiar_cuit(raw)

    if not valor:
        return jsonify({"error": "Ingresá un CUIT para buscar."}), 400
    if not valor.isdigit() or len(valor) != 11:
        return jsonify({"error": f"CUIT inválido. Debe tener 11 dígitos. Se recibió: '{raw}'"}), 400

    try:
        conn, tipo = get_connection()
        cursor = conn.cursor()
        all_fields = READ_FIELDS + EDIT_FIELDS
        fields_sql = ", ".join(all_fields)
        cursor.execute(f"SELECT {fields_sql} FROM {TABLE_NAME} WHERE {SEARCH_FIELD} = ?", (valor,))
        row = cursor.fetchone()
        conn.close()

        if row is None:
            return jsonify({"error": f"No se encontró ningún registro con CUIT {valor}."}), 404

        result = dict(zip(all_fields, row))
        for k, v in result.items():
            if v is not None and not isinstance(v, (str, int, float, bool)):
                result[k] = str(v)
        return jsonify({"data": result})
    except Exception:
        log.exception("Error en /buscar")
        return jsonify({"error": "Error interno al consultar la base de datos."}), 500


@app.route("/actualizar", methods=["POST"])
def actualizar():
    payload      = request.json
    search_value = limpiar_cuit(payload.get("search_value", ""))
    updates      = payload.get("updates", {})
    usuario      = payload.get("usuario", "").strip()

    if not search_value:
        return jsonify({"error": "CUIT no recibido."}), 400

    updates_filtrados = {k: v for k, v in updates.items() if k in EDIT_FIELDS}
    use_getdate = "Fech_created" in updates_filtrados and updates_filtrados["Fech_created"] == "__GETDATE__"
    if use_getdate:
        del updates_filtrados["Fech_created"]

    if "UPass" in updates_filtrados:
        upass_val = (updates_filtrados["UPass"] or "").strip()
        if upass_val and not UPASS_RE.match(upass_val):
            return jsonify({"error": "UPass inválido. Debe ser exactamente 8 caracteres alfanuméricos."}), 400
        updates_filtrados["UPass"] = upass_val or None

    if "Email" in updates_filtrados:
        email_val = (updates_filtrados["Email"] or "").strip()
        if email_val and not EMAIL_RE.match(email_val):
            return jsonify({"error": "Email inválido."}), 400
        updates_filtrados["Email"] = email_val or None

    updates_filtrados["auditoria"] = usuario if usuario else "sistema"

    try:
        conn, tipo = get_connection()
        cursor = conn.cursor()

        set_parts = []
        values    = []

        if use_getdate:
            fecha_fn = "DATE('now')" if tipo == "sqlite" else "CAST(GETDATE() AS DATE)"
            set_parts.append(f"Fech_created = {fecha_fn}")

        for k, v in updates_filtrados.items():
            set_parts.append(f"{k} = ?")
            values.append(v)

        values.append(search_value)
        set_clause = ", ".join(set_parts)
        cursor.execute(f"UPDATE {TABLE_NAME} SET {set_clause} WHERE {SEARCH_FIELD} = ?", values)
        conn.commit()
        affected = cursor.rowcount
        conn.close()

        if affected == 0:
            return jsonify({"error": "No se actualizó ningún registro."}), 404
        return jsonify({"ok": True, "mensaje": "Registro actualizado correctamente."})
    except Exception:
        log.exception("Error en /actualizar")
        return jsonify({"error": "Error interno al actualizar el registro."}), 500



@app.route("/generar_doc", methods=["POST"])
def generar_doc():
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import io

    payload = request.json
    cuit    = payload.get("cuit", "").strip()
    upass   = payload.get("upass", "").strip()

    if not cuit or not upass:
        return jsonify({"error": "Faltan datos para generar el documento."}), 400

    doc = Document()

    # Ajustar márgenes
    from docx.shared import Inches
    for section in doc.sections:
        section.top_margin    = Inches(1.5)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.5)
        section.right_margin  = Inches(1.5)

    def linea_grande(label, valor=None):
        p = doc.add_paragraph()
        r1 = p.add_run(label)
        r1.font.size = Pt(24)
        if valor:
            r2 = p.add_run(" " + valor)
            r2.font.size = Pt(24)
        return p

    def linea_chica(texto):
        p = doc.add_paragraph()
        for palabra in texto.split(" "):
            r = p.add_run(palabra + " ")
            r.font.size = Pt(15)
        return p

    # Espacios iniciales
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    linea_grande("PAGINA WEB:", "http://tesoreria.mecontuc.gob.ar/")
    linea_grande("USUARIO:",    cuit)
    linea_grande("CONTRASEÑA:", upass)

    doc.add_paragraph()
    doc.add_paragraph()

    linea_chica("Recuerde habilitar ventanas emergentes en el navegador.")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    from flask import send_file
    return send_file(
        buf,
        as_attachment=True,
        download_name=f"credenciales_{cuit}.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


if __name__ == "__main__":
    if MODO == "sqlite":
        init_sqlite()
    app.run(host="0.0.0.0", port=5000, debug=False)
